from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(r"C:\Users\唐乐\Desktop\个人网站")
ASSETS = ROOT / "assets"
DOCX_OUT = ASSETS / "tangle-resume.docx"
PDF_OUT = ASSETS / "tangle-resume.pdf"


DATA = {
    "name": "唐乐",
    "target": "软件开发工程师 / Java 开发 / Python 开发 / 微信小程序开发",
    "contact": "18064871268 | 2061790875@qq.com | 云南昆明",
    "badges": "上海应用技术大学 · 软件工程本科 · 2023.10-2027.06 | CET-6 | 入党积极分子",
    "education": "上海应用技术大学　软件工程　本科　2023.10 - 2027.06",
    "skills": [
        "编程语言：Java、Python、C++、JavaScript、HTML/CSS",
        "数据库与后端：MySQL、数据库表设计、CRUD、基础查询优化",
        "小程序与前端：微信小程序原生开发、WXML、WXSS、页面交互实现",
        "工程工具：GitHub、微信开发者工具、微信云开发、Node.js Test Runner",
    ],
    "project": {
        "name": "星途自律舱 - 星球养成式自律打卡微信小程序",
        "time": "2026.06",
        "bullets": [
            "独立完成需求拆解、页面设计与核心功能开发，围绕运动、饮食、学习、工作、计划、睡眠构建六维打卡体系。",
            "设计“目标设定 - 自定义计划 - 打卡反馈 - 星球成长”体验闭环，实现星光奖励、宝箱、成就墙、生态区入口和数据统计模块。",
            "接入微信云开发，完成登录、打卡提交、统计读取、资料更新、数据导出、云资源链接等云函数与本地数据兜底逻辑。",
            "处理小程序包体限制与资源加载问题，将大图、视频拆分至云存储，配合本地压缩素材、缓存策略和错误降级提示提升稳定性。",
            "使用 Node.js Test Runner 覆盖奖励计算、目标进度、成就解锁和统计聚合等核心逻辑，减少手动调试成本。",
        ],
    },
    "next_project": "下一项目经历预留：后续完成后可补充项目名称、技术栈、职责与成果。",
    "practice": [
        "校园运动会志愿服务：负责检录、赛道引导、物资分发和成绩登记协助，提升现场执行与沟通能力。",
        "社区电脑维修公益服务：为居民提供系统调试、软件安装、网络故障排查等技术支持，增强动手排障能力。",
    ],
    "summary": "软件工程专业基础较扎实，具备 Java、Python、MySQL 和微信小程序开发实践；能独立推进小型项目从需求拆解到开发测试落地。学习能力强，责任心较好，希望在真实项目中继续积累工程经验。",
}


def set_font(run, size=8.5, bold=False, color="1F2937"):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, 11, True, "1D4ED8")
    p_pr = p._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "D7DEE8")
    border.append(bottom)
    p_pr.append(border)


def add_text(doc, text, size=8.5, bold=False, color="1F2937", after=1.0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(text)
    set_font(run, size, bold, color)


def add_bullet(doc, text, size=8.2, color="1F2937"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.38)
    p.paragraph_format.first_line_indent = Cm(-0.18)
    p.paragraph_format.space_after = Pt(0.8)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run("- " + text)
    set_font(run, size, False, color)


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.05)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.25)
    section.right_margin = Cm(1.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(8.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(DATA["name"])
    set_font(run, 20, True, "111827")

    for text, size, color in [
        (DATA["target"], 9.2, "374151"),
        (DATA["contact"], 8.5, "4B5563"),
        (DATA["badges"], 8.5, "B45309"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        set_font(run, size, False, color)

    add_heading(doc, "教育背景")
    add_text(doc, DATA["education"], 8.8, True)

    add_heading(doc, "技术能力")
    table = doc.add_table(rows=2, cols=2)
    table.autofit = True
    for i, skill in enumerate(DATA["skills"]):
        cell = table.cell(i // 2, i % 2)
        cell.text = ""
        tc_pr = cell._tc.get_or_add_tcPr()
        shade = OxmlElement("w:shd")
        shade.set(qn("w:fill"), "F7FAFC")
        tc_pr.append(shade)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(skill)
        set_font(run, 8.0)

    add_heading(doc, "项目经历")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    left = p.add_run(DATA["project"]["name"])
    set_font(left, 9.3, True, "111827")
    right = p.add_run("    " + DATA["project"]["time"])
    set_font(right, 8.4, False, "6B7280")
    for bullet in DATA["project"]["bullets"]:
        add_bullet(doc, bullet, 8.1)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(7)
    run = p.add_run(DATA["next_project"])
    set_font(run, 8.2, False, "9CA3AF")

    add_heading(doc, "实践经历")
    for item in DATA["practice"]:
        add_bullet(doc, item, 7.9)

    add_heading(doc, "自我评价")
    add_bullet(doc, DATA["summary"], 8.0)

    ASSETS.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_OUT)


def pstyle(name, size, leading, color="#1f2937", bold=False, align=0):
    return ParagraphStyle(
        name,
        fontName="MSYHBD" if bold else "MSYH",
        fontSize=size,
        leading=leading,
        textColor=colors.HexColor(color),
        alignment=align,
        spaceAfter=2,
    )


def build_pdf():
    pdfmetrics.registerFont(TTFont("MSYH", r"C:\Windows\Fonts\msyh.ttc"))
    pdfmetrics.registerFont(TTFont("MSYHBD", r"C:\Windows\Fonts\msyhbd.ttc"))

    doc = SimpleDocTemplate(
        str(PDF_OUT),
        pagesize=A4,
        rightMargin=13 * mm,
        leftMargin=13 * mm,
        topMargin=10 * mm,
        bottomMargin=9 * mm,
    )
    title = pstyle("title", 18, 22, "#111827", True, 1)
    center = pstyle("center", 8.2, 10, "#4b5563", False, 1)
    badge = pstyle("badge", 8.2, 10, "#b45309", False, 1)
    h = pstyle("h", 10.5, 13, "#1d4ed8", True)
    body = pstyle("body", 7.7, 9.4)
    small = pstyle("small", 7.45, 9.1, "#374151")
    proj = pstyle("proj", 8.7, 10.5, "#111827", True)
    muted = pstyle("muted", 7.8, 9.6, "#9ca3af")

    story = [
        Paragraph(DATA["name"], title),
        Paragraph(DATA["target"], center),
        Paragraph(DATA["contact"], center),
        Paragraph(DATA["badges"], badge),
        Spacer(1, 4),
    ]

    def section(text):
        story.append(Paragraph(text, h))
        story.append(Table([[""]], colWidths=[184 * mm], rowHeights=[1], style=TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#d7dee8"))
        ])))
        story.append(Spacer(1, 2))

    section("教育背景")
    story.append(Paragraph(DATA["education"], body))

    section("技术能力")
    skill_rows = [
        [Paragraph(DATA["skills"][0], small), Paragraph(DATA["skills"][1], small)],
        [Paragraph(DATA["skills"][2], small), Paragraph(DATA["skills"][3], small)],
    ]
    skill_table = Table(skill_rows, colWidths=[92 * mm, 92 * mm])
    skill_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f7fafc")),
        ("BOX", (0, 0), (-1, -1), 0.3, colors.HexColor("#e5e7eb")),
        ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#e5e7eb")),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    story.append(skill_table)

    section("项目经历")
    story.append(Paragraph(f'{DATA["project"]["name"]}　<span color="#6b7280">{DATA["project"]["time"]}</span>', proj))
    for bullet in DATA["project"]["bullets"]:
        story.append(Paragraph("- " + bullet, body))
    story.append(Spacer(1, 5))
    story.append(Paragraph(DATA["next_project"], muted))
    story.append(Spacer(1, 7))

    section("实践经历")
    for item in DATA["practice"]:
        story.append(Paragraph("- " + item, body))

    section("自我评价")
    story.append(Paragraph("- " + DATA["summary"], body))
    doc.build(story)


if __name__ == "__main__":
    build_docx()
    build_pdf()
    print(DOCX_OUT)
    print(PDF_OUT)
